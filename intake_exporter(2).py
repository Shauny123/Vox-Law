from docx import Document
from datetime import datetime
import os
import subprocess
import requests
import json
import socket
from rag_utils import export_and_upload_document, log_client_edit_event, queue_for_learning_pipeline
from fastapi import FastAPI
import uvicorn
from transformers import pipeline
import openai
import anthropic

app = FastAPI()

# Whisper + Flamingo fallback
asr_pipeline = pipeline("automatic-speech-recognition", model="openai/whisper-large")
def fallback_transcribe(audio_path):
    try:
        return asr_pipeline(audio_path)["text"]
    except:
        # Fallback to NVIDIA Flamingo server
        return requests.post("http://localhost:5001/flamingo/transcribe", files={"file": open(audio_path, "rb")}).text

# GPT-4 / Claude / Kimi fallback
openai.api_key = os.getenv("OPENAI_API_KEY")
claude_client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
def gpt_fallback(prompt):
    try:
        return openai.ChatCompletion.create(model="gpt-4", messages=[{"role": "user", "content": prompt}])['choices'][0]['message']['content']
    except:
        return claude_client.messages.create(
            model="claude-3-opus-20240229", 
            max_tokens=800,
            messages=[{"role": "user", "content": prompt}]
        ).content


def sanitize_filename(name):
    return name.replace(" ", "_").replace("/", "_").lower()

def add_disclaimer_to_doc(doc):
    doc.add_paragraph("\n---\n")
    doc.add_paragraph("⚠️ Disclaimer: This document was generated using artificial intelligence (AI). It is not legal advice and must be reviewed and approved by a licensed attorney before being used for any legal purpose. Submitting AI-generated documents directly to a court may result in sanctions. Paths Apart LLC assumes no responsibility for the content, interpretations, or outcomes resulting from this document. This document is for informational purposes only and intended solely to assist in seeking legal representation.", style='Intense Quote')

def find_open_port(start_port=8000, max_tries=50):
    for port in range(start_port, start_port + max_tries):
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            if s.connect_ex(('localhost', port)) != 0:
                return port
    raise RuntimeError("No open ports found.")

def auto_export_and_upload(doc_data, bucket="your-intake-bucket"):
    now = datetime.utcnow().strftime("%Y%m%d-%H%M")
    client_name = sanitize_filename(doc_data.get("name", "anonymous"))
    case_type = sanitize_filename(doc_data.get("case_type", "general"))

    base_filename = f"{client_name}_{case_type}_{now}"
    docx_file = f"/tmp/{base_filename}.docx"
    pdf_file = f"/tmp/{base_filename}.pdf"
    remote_docx = f"intakes/{case_type}/{base_filename}.docx"
    remote_pdf = f"intakes/{case_type}/{base_filename}.pdf"

    doc = Document()
    doc.add_heading("Client Legal Intake Summary", 0)
    for k, v in doc_data.items():
        doc.add_paragraph(f"{k.capitalize()}: {v}")
    add_disclaimer_to_doc(doc)
    doc.save(docx_file)

    try:
        subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", docx_file, "--outdir", "/tmp"], check=True)
    except Exception as e:
        print(f"[ERROR] PDF conversion failed: {e}")

    export_and_upload_document(docx_file, bucket, remote_docx)
    export_and_upload_document(pdf_file, bucket, remote_pdf)

    return remote_docx, remote_pdf

def trigger_n8n_webhook(gcs_docx, gcs_pdf, intake_data, webhook_url="https://n8n.yourdomain/webhook/intake"):
    payload = {
        "docx_url": f"https://storage.googleapis.com/your-intake-bucket/{gcs_docx}",
        "pdf_url": f"https://storage.googleapis.com/your-intake-bucket/{gcs_pdf}",
        "client_email": intake_data.get("email"),
        "intake_data": intake_data
    }
    try:
        r = requests.post(webhook_url, json=payload)
        r.raise_for_status()
        print(f"n8n webhook triggered for {intake_data.get('name')}")
    except Exception as e:
        print(f"[ERROR] Failed to notify webhook: {e}")

def track_client_pause_event(case_id, timestamp, transcription_segment):
    print(f"Pause logged: case={case_id}, time={timestamp}, text={transcription_segment}")
    log_client_edit_event(case_id, timestamp, transcription_segment)

def queue_for_learning(case_id, edit_type, raw_text, corrected_text):
    queue_for_learning_pipeline(case_id, edit_type, raw_text, corrected_text)
    print(f"Queued edit for learning: {edit_type}: '{raw_text}' -> '{corrected_text}'")

def upload_locked_document(file_path, client_id, filename, bucket="your-intake-bucket"):
    remote_path = f"locked_sources/{client_id}/{filename}"
    metadata_path = f"/tmp/{filename}.meta.json"
    metadata = {
        "locked": True,
        "uploaded_by": client_id,
        "timestamp": datetime.utcnow().isoformat()
    }
    with open(metadata_path, 'w') as meta:
        json.dump(metadata, meta)
    export_and_upload_document(file_path, bucket, remote_path)
    export_and_upload_document(metadata_path, bucket, f"locked_sources/{client_id}/{filename}.meta.json")
    print(f"Locked upload completed: {remote_path}")

# Self-healing runtime (placeholder for retry + alert logic)
def ensure_auto_launch_service(service_func, max_retries=3):
    for attempt in range(max_retries):
        try:
            port = find_open_port()
            print(f"Launching service on port {port}...")
            service_func(port)
            return
        except Exception as e:
            print(f"[RETRY {attempt+1}] Service failed to start: {e}")
    print("[FATAL] Service could not be launched after retries.")

def launch_fastapi_service(port):
    uvicorn.run("intake_exporter:app", host="0.0.0.0", port=port, reload=True)

if __name__ == "__main__":
    ensure_auto_launch_service(launch_fastapi_service)
